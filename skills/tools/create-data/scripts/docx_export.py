"""DOCX export with randomized document style and structure variations.

Variations per export:
- Style: formal report, listing, memo
- Font: Calibri, Arial, Times New Roman, Cambria
- Font size: 10, 11, or 12pt
- Line spacing: 1.0, 1.15, 1.5
- Accent color: 5 distinct palettes
- Optional: table vs paragraph presentation, table of contents, page breaks
"""

from __future__ import annotations

import json
import random
from pathlib import Path
from typing import Any

from helpers import truncate

_FONTS = ["Calibri", "Arial", "Times New Roman", "Cambria"]


def export_to_docx(data: list[dict[str, Any]], path: str | Path) -> Path:
    """Export data to DOCX with randomized document style and structure."""
    from docx import Document
    from docx.shared import Pt, RGBColor

    dest = Path(path).with_suffix(".docx")
    dest.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    style_choice = random.choice(["report", "listing", "memo"])
    font_name = random.choice(_FONTS)
    font_size = random.choice([10, 11, 12])
    use_tables = random.choice([True, False])
    use_page_breaks = random.choice([True, False]) and len(data) <= 20
    use_toc_header = random.choice([True, False])
    line_spacing = random.choice([1.0, 1.15, 1.5])

    style = doc.styles["Normal"]
    style.font.name = font_name
    style.font.size = Pt(font_size)
    style.paragraph_format.line_spacing = line_spacing

    accent_colors = [
        RGBColor(0x1F, 0x4E, 0x79),
        RGBColor(0x54, 0x82, 0x35),
        RGBColor(0xBF, 0x8F, 0x00),
        RGBColor(0xC0, 0x00, 0x00),
        RGBColor(0x70, 0x30, 0xA0),
    ]
    accent = random.choice(accent_colors)

    if style_choice == "report":
        _report_style(doc, data, font_name, font_size, accent, use_tables, use_page_breaks)
    elif style_choice == "listing":
        _listing_style(doc, data, font_name, font_size, accent, use_toc_header)
    else:
        _memo_style(doc, data, font_name, font_size, accent)

    doc.save(str(dest))
    return dest


def _report_style(doc, data, font_name, font_size, accent, use_tables, use_page_breaks):
    from docx.shared import Pt

    title = doc.add_heading("Dataset Export Report", level=0)
    for run in title.runs:
        run.font.color.rgb = accent

    doc.add_paragraph(f"Total entries: {len(data)}")

    scores = [
        (e.get("metadata") or {}).get("quality_score")
        for e in data
        if (e.get("metadata") or {}).get("quality_score") is not None
    ]
    if scores:
        doc.add_paragraph(f"Average quality score: {sum(scores) / len(scores):.1f}")

    doc.add_page_break()

    for idx, entry in enumerate(data, 1):
        h = doc.add_heading(f"Entry {idx}", level=2)
        for run in h.runs:
            run.font.color.rgb = accent

        meta = entry.get("metadata") or {}

        if use_tables and meta:
            table = doc.add_table(rows=1, cols=2)
            table.style = "Light Grid Accent 1"
            hdr = table.rows[0].cells
            hdr[0].text = "Property"
            hdr[1].text = "Value"
            for k, v in meta.items():
                row = table.add_row().cells
                row[0].text = k.replace("_", " ").title()
                row[1].text = json.dumps(v) if isinstance(v, (list, dict)) else str(v)
            doc.add_paragraph()

        doc.add_heading("Input", level=3)
        doc.add_paragraph(entry.get("input", ""))

        doc.add_heading("Output", level=3)
        p = doc.add_paragraph(entry.get("output", ""))
        for run in p.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(font_size - 1)

        if use_page_breaks and idx < len(data):
            doc.add_page_break()
        else:
            doc.add_paragraph("---")


def _listing_style(doc, data, font_name, font_size, accent, use_toc_header):
    from docx.shared import Pt

    title = doc.add_heading("Generated Dataset", level=1)
    for run in title.runs:
        run.font.color.rgb = accent

    if use_toc_header:
        doc.add_heading("Contents", level=2)
        for idx, entry in enumerate(data, 1):
            req = (entry.get("metadata") or {}).get("request_type", "Entry")
            doc.add_paragraph(f"{idx}. {req}", style="List Number")
        doc.add_page_break()

    for idx, entry in enumerate(data, 1):
        meta = entry.get("metadata") or {}
        req = meta.get("request_type", "Entry")
        score = meta.get("quality_score", "")

        h = doc.add_heading(f"{idx}. {req}", level=2)
        for run in h.runs:
            run.font.color.rgb = accent

        if score:
            doc.add_paragraph(f"Quality Score: {score}")

        rules = entry.get("transformation_rules", [])
        if rules:
            doc.add_paragraph("Rules:", style="List Bullet")
            for r in rules:
                doc.add_paragraph(r, style="List Bullet 2")

        p = doc.add_paragraph()
        run = p.add_run("Input: ")
        run.bold = True
        p.add_run(entry.get("input", ""))

        p = doc.add_paragraph()
        run = p.add_run("Output: ")
        run.bold = True
        out_run = p.add_run(entry.get("output", ""))
        out_run.font.name = "Courier New"
        out_run.font.size = Pt(font_size - 1)

        doc.add_paragraph()


def _memo_style(doc, data, font_name, font_size, accent):
    title = doc.add_heading("MEMO: Dataset Export", level=1)
    for run in title.runs:
        run.font.color.rgb = accent

    p = doc.add_paragraph()
    p.add_run("To: ").bold = True
    p.add_run("Data Review Team")
    p = doc.add_paragraph()
    p.add_run("Subject: ").bold = True
    p.add_run(f"Generated Dataset ({len(data)} entries)")
    p = doc.add_paragraph()
    p.add_run("Summary: ").bold = True
    scores = [
        (e.get("metadata") or {}).get("quality_score")
        for e in data
        if (e.get("metadata") or {}).get("quality_score") is not None
    ]
    if scores:
        p.add_run(f"Avg score {sum(scores) / len(scores):.1f} across {len(data)} entries")
    else:
        p.add_run(f"{len(data)} entries generated")

    doc.add_paragraph("---")

    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Shading Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "#"
    hdr[1].text = "Type"
    hdr[2].text = "Input (Preview)"
    hdr[3].text = "Score"

    for idx, entry in enumerate(data, 1):
        meta = entry.get("metadata") or {}
        row = table.add_row().cells
        row[0].text = str(idx)
        row[1].text = meta.get("request_type", "N/A")
        row[2].text = truncate(entry.get("input", ""), 80)
        row[3].text = str(meta.get("quality_score", ""))

    doc.add_paragraph()
    doc.add_heading("Detailed Entries", level=2)

    for idx, entry in enumerate(data, 1):
        doc.add_heading(f"Entry {idx}", level=3)
        doc.add_paragraph(entry.get("input", ""))
        p = doc.add_paragraph()
        run = p.add_run(entry.get("output", ""))
        run.italic = True
        doc.add_paragraph()
